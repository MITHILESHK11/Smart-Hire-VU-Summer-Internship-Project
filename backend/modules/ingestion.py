import io
import logging
import fitz  # PyMuPDF
import docx

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts raw text from PDF bytes using PyMuPDF (fitz)."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to parse PDF document: {e}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts raw text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to parse Word document: {e}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Detects format from filename extension and extracts raw text."""
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif lower_filename.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format: {filename}. Supported formats are PDF, DOCX, and TXT.")
